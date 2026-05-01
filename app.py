import streamlit as st
import pandas as pd
import xmlrpc.client
import base64
import os
import io
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
ODOO_URL     = os.environ.get("ODOO_URL",     "https://inovues.odoo.com")
ODOO_DB      = os.environ.get("ODOO_DB",      "inovues")
ODOO_USER    = os.environ.get("ODOO_USER",    "sketterer@inovues.com")
ODOO_API_KEY = os.environ.get("ODOO_API_KEY", "")

SHIP_TO_DEFAULT = "Momentum Glass, LLC\nAttn: INOVUES, INC.\n25825 Aldine Westfield Rd.\nSpring, TX 77373\n281.809.2830"

sq_inches_to_sq_feet = 1 / 144

st.set_page_config(
    page_title="INOVUES PO Generator",
    page_icon="📝",
    layout="wide"
)

# ─────────────────────────────────────────────────────────────
# ODOO HELPERS
# ─────────────────────────────────────────────────────────────
def get_odoo_connection():
    common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
    uid = common.authenticate(ODOO_DB, ODOO_USER, ODOO_API_KEY, {})
    if not uid:
        raise Exception("Odoo authentication failed — check ODOO_API_KEY.")
    models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
    return uid, models

def odoo_call(models, uid, model, method, args, kwargs={}):
    return models.execute_kw(ODOO_DB, uid, ODOO_API_KEY, model, method, args, kwargs)

@st.cache_data(ttl=120, show_spinner="Loading projects from Odoo...")
def fetch_projects():
    try:
        uid, models = get_odoo_connection()
        projects = odoo_call(models, uid, "project.project", "search_read",
            [[("active", "=", True)]],
            {"fields": ["id", "name"], "order": "name asc", "limit": 200})
        return {p["name"]: p["id"] for p in projects}, None
    except Exception as e:
        return {}, str(e)

@st.cache_data(ttl=60, show_spinner="Fetching files from Odoo...")
def fetch_project_attachments(project_id, po_type_key="Glass"):
    """Get attachments from the relevant tasks in this project.

    For Glass POs, looks for SWR Cutlist tasks (which produce *_SWR_Glass_*.xlsx).
    For Aluminium POs, looks for Cutting Optimization tasks (which produce cutting_list_*.xlsx).
    """
    try:
        uid, models = get_odoo_connection()

        # Choose task name keyword based on PO type
        task_keyword = "Cutting Optimization" if po_type_key == "Aluminium" else "SWR"

        # Find tasks in this project matching the keyword
        tasks = odoo_call(models, uid, "project.task", "search_read",
            [[("project_id", "=", project_id), ("name", "ilike", task_keyword)]],
            {"fields": ["id", "name", "create_date"], "order": "create_date desc", "limit": 10})

        if not tasks:
            # Fall back to any task in the project
            tasks = odoo_call(models, uid, "project.task", "search_read",
                [[("project_id", "=", project_id)]],
                {"fields": ["id", "name", "create_date"], "order": "create_date desc", "limit": 20})

        if not tasks:
            return [], None

        # Get attachments from all found tasks
        task_ids = [t["id"] for t in tasks]
        attachments = odoo_call(models, uid, "ir.attachment", "search_read",
            [[("res_model", "=", "project.task"), ("res_id", "in", task_ids),
               ("name", "ilike", ".xlsx")]],
            {"fields": ["id", "name", "res_id", "datas", "create_date"],
             "order": "create_date desc"})

        # Add task name context
        task_map = {t["id"]: t["name"] for t in tasks}
        for a in attachments:
            a["task_name"] = task_map.get(a["res_id"], "Unknown task")

        return attachments, None
    except Exception as e:
        return [], str(e)

@st.cache_data(ttl=30, show_spinner="Loading vendors from Odoo...")
def fetch_vendors():
    try:
        uid, models = get_odoo_connection()
        vendors = odoo_call(models, uid, "res.partner", "search_read",
            [[("supplier_rank", ">", 0)]],
            {"fields": ["id", "name", "email", "phone", "street", "street2",
                        "city", "state_id", "zip", "country_id", "child_ids"],
             "order": "name asc", "limit": 200})
        if not vendors:
            vendors = odoo_call(models, uid, "res.partner", "search_read",
                [[("is_company", "=", True)]],
                {"fields": ["id", "name", "email", "phone", "street", "street2",
                            "city", "state_id", "zip", "country_id", "child_ids"],
                 "order": "name asc", "limit": 200})
        # Enrich with address + contact
        for v in vendors:
            addr_parts = [p for p in [v.get("street"), v.get("street2")] if p]
            city_line = ", ".join(p for p in [
                v.get("city"),
                v.get("state_id", [False, ""])[1] if isinstance(v.get("state_id"), list) else "",
                v.get("zip")
            ] if p)
            if city_line:
                addr_parts.append(city_line)
            country = v.get("country_id", [False, ""])[1] if isinstance(v.get("country_id"), list) else ""
            if country:
                addr_parts.append(country)
            v["full_address"] = "\n".join(addr_parts)

            # Pick contact: prefer child tagged 'Orders', else first child.
            # The 'Orders' tag (res.partner.category) flags the primary ordering
            # contact for a vendor — set it on whichever child should appear in POs.
            v["contact_name"] = ""
            if v.get("child_ids"):
                primary_ids = odoo_call(models, uid, "res.partner", "search",
                    [[("id", "in", v["child_ids"]),
                      ("category_id.name", "=", "Orders")]],
                    {"limit": 1})
                target_ids = primary_ids if primary_ids else v["child_ids"][:1]
                contacts = odoo_call(models, uid, "res.partner", "read",
                    [target_ids], {"fields": ["name", "email"]})
                if contacts:
                    v["contact_name"] = contacts[0].get("name", "")
                    if not v.get("email"):
                        v["email"] = contacts[0].get("email", "")
        return {v["name"]: v for v in vendors}, None
    except Exception as e:
        return {}, str(e)

@st.cache_data(ttl=300, show_spinner="Loading payment terms...")
def fetch_payment_terms():
    """Returns ({display_name: id}, error). Cached 5 min — these change rarely."""
    try:
        uid, models = get_odoo_connection()
        terms = odoo_call(models, uid, "account.payment.term", "search_read",
            [[("active", "=", True)]],
            {"fields": ["id", "name"], "order": "sequence asc, id asc"})
        return {t["name"]: t["id"] for t in terms}, None
    except Exception as e:
        return {}, str(e)

@st.cache_data(ttl=300, show_spinner="Loading incoterms...")
def fetch_incoterms():
    """Returns ({code: id}, error). Cached 5 min."""
    try:
        uid, models = get_odoo_connection()
        terms = odoo_call(models, uid, "account.incoterms", "search_read",
            [[("active", "=", True)]],
            {"fields": ["id", "code", "name"], "order": "code asc"})
        # Display "DDP — DELIVERED DUTY PAID" but key by code so user picks easily
        return {f"{t['code']} — {t['name']}": t["id"] for t in terms}, None
    except Exception as e:
        return {}, str(e)

@st.cache_data(ttl=300, show_spinner="Loading users...")
def fetch_users():
    """Returns ({display_name: id}, error). Internal users only.
    Used for fuzzy-matching the free-text requisitioner field to a user_id."""
    try:
        uid, models = get_odoo_connection()
        users = odoo_call(models, uid, "res.users", "search_read",
            [[("share", "=", False), ("active", "=", True)]],
            {"fields": ["id", "name", "login"], "order": "name asc"})
        # Build a map keyed by both name AND login email for fuzzy matching
        user_map = {}
        for u in users:
            user_map[u["name"].lower()] = u["id"]
            user_map[u["login"].lower()] = u["id"]
            # Also map first name only for casual matches like "Stephan"
            first = u["name"].split()[0].lower() if u["name"] else ""
            if first and first not in user_map:
                user_map[first] = u["id"]
        return user_map, None
    except Exception as e:
        return {}, str(e)

# ─────────────────────────────────────────────────────────────
# FILE PARSERS
# ─────────────────────────────────────────────────────────────
def parse_glass_file(file_bytes):
    """Parse SWR_Glass xlsx — returns list of glass line items."""
    try:
        # Glass file has header info in first 12 rows, data starts at row 12 (0-indexed)
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="Glass", header=12)
        # If Tag column not found, try finding it dynamically
        if "Tag" not in df.columns:
            raw = pd.read_excel(io.BytesIO(file_bytes), sheet_name="Glass", header=None)
            for i, row in raw.iterrows():
                if "Tag" in str(row.values):
                    df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="Glass", header=i)
                    break
        # Drop totals row
        df = df[df["Tag"].astype(str).str.strip() != "Totals"].copy()
        df = df.dropna(subset=["Tag"])
        lines = []
        for _, row in df.iterrows():
            try:
                area_each  = float(row.get("Area Each (ft²)", 0) or 0)
                qty        = int(float(row.get("Qty", 0) or 0))
                area_total = float(row.get("Area Total (ft²)", 0) or 0)
                w_16 = str(row.get("Glass Width (1/16)", "")).strip()
                h_16 = str(row.get("Glass Height (1/16)", "")).strip()
                size_str = f'{w_16}" x {h_16}"' if w_16 and h_16 else ""
                if qty > 0:
                    lines.append({
                        "tag":        str(row.get("Tag", "")),
                        "size_str":   size_str,
                        "area_each":  round(area_each, 2),
                        "qty":        qty,
                        "area_total": round(area_total, 2),
                        "description": "",
                    })
            except Exception:
                continue
        return lines, None
    except Exception as e:
        return [], str(e)

def parse_aggcutonly_file(file_bytes):
    """Parse SWR_AggCutOnly xlsx — returns list of aluminium cut items."""
    try:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="AggCutOnly")
        if "Finished Length in" not in df.columns:
            # Try to find header row
            for i, row in df.iterrows():
                if "Finished Length in" in str(row.values):
                    df = pd.read_excel(io.BytesIO(file_bytes), sheet_name="AggCutOnly",
                                       header=i)
                    break
        df = df.dropna(subset=["Finished Length in"])
        lines = []
        for _, row in df.iterrows():
            try:
                length   = float(row.get("Finished Length in", 0) or 0)
                total_qty = int(float(row.get("Total QTY", 0) or 0))
                part_num  = str(row.get("Part #", "")).strip()
                if length > 0 and total_qty > 0:
                    lines.append({
                        "length_in":   round(length, 3),
                        "length_ft":   round(length / 12, 3),
                        "part_number": part_num,
                        "qty":         total_qty,
                        "description": f"{part_num} — {length:.3f}\" extrusion",
                    })
            except Exception:
                continue
        return lines, None
    except Exception as e:
        return [], str(e)

def parse_optimizer_file(file_bytes):
    """Parse cutting_list optimizer Excel — returns line items grouped by (Part Number, Stock Length)."""
    try:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=None)

        # Find the DETAILED CUT LIST header row
        detail_row = None
        for i, row in df.iterrows():
            if str(row.iloc[0]).strip() == "DETAILED CUT LIST":
                detail_row = i
                break
        if detail_row is None:
            return [], "Could not find 'DETAILED CUT LIST' section in optimizer file."

        # The column header row is right after
        col_row = detail_row + 1
        df_detail = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=col_row)

        # Keep only rows that have Stock Length (ft) — these are the stock piece rows
        df_detail = df_detail.dropna(subset=["Stock Length (ft)"])
        df_detail = df_detail[pd.to_numeric(df_detail["Stock Length (ft)"], errors="coerce").notna()]

        # Group by (Part Number, Stock Length (ft)) and count rows = number of stock pieces
        df_detail["Stock Length (ft)"] = pd.to_numeric(df_detail["Stock Length (ft)"], errors="coerce")
        df_detail["Part Number"] = df_detail["Part Number"].fillna("").astype(str).str.strip()

        grouped = (
            df_detail.groupby(["Part Number", "Stock Length (ft)"])
            .size()
            .reset_index(name="qty")
        )

        lines = []
        for _, row in grouped.iterrows():
            length_ft = float(row["Stock Length (ft)"])
            length_in = round(length_ft * 12, 3)
            lines.append({
                "profile":     str(row["Part Number"]),
                "length_in":   length_in,
                "length_ft":   round(length_ft, 3),
                "qty":         int(row["qty"]),
                # PO-specific fields — user fills in app
                "alloy":       "6063-T6",
                "finish":      "",
                "lead_time":   "3-4 weeks",
                "tooling":     "",
                "die_setup":   "",
                "unit_price":  0.0,
                "um":          "Pc",
            })
        return lines, None
    except Exception as e:
        return [], str(e)


# ─────────────────────────────────────────────────────────────
# ALUMINUM PO DOCX GENERATOR
# ─────────────────────────────────────────────────────────────
def generate_aluminum_po_docx(
    vendor_name, vendor_contact, vendor_email, vendor_address,
    ship_to_lines, job_number, job_location,
    po_date, po_number, requisitioner,
    shipped_via, fob_point, terms,
    al_lines,          # list of dicts with profile, length_in, qty, alloy, finish, lead_time, tooling, die_setup, unit_price, um
    packaging_note="Fully Corrugated Bundles, Paper Layer Separation",
    logo_path=None,
):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin   = Cm(1.27)
        section.right_margin  = Cm(1.27)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    subtotal   = sum(l["unit_price"] * l["qty"] for l in al_lines)
    grand_total = subtotal  # shipping/tax added below if non-zero

    table = doc.add_table(rows=0, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Columns: Item# | Profile# | Alloy | Finish | Length(in) | Lead Time | Pcs | Unit Price | Total
    col_widths = [Cm(0.8), Cm(1.8), Cm(1.8), Cm(2.8), Cm(1.6), Cm(2.2), Cm(1.2), Cm(2.0), Cm(2.2)]
    for i, w in enumerate(col_widths):
        table.columns[i].width = w

    # ── Logo + PURCHASE ORDER ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[3])
    if logo_path and os.path.exists(logo_path):
        c.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))
    c = row.cells[4]; c.merge(row.cells[8])
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("PURCHASE ORDER")
    run.bold = True; run.font.size = Pt(16); run.font.small_caps = True

    # ── TO / SHIP TO / PO NUMBER ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[2])
    p = c.paragraphs[0]; run = p.add_run("TO:"); run.bold = True; run.font.small_caps = True
    if vendor_contact: c.add_paragraph().add_run(vendor_contact)
    c.add_paragraph().add_run(vendor_name)
    if vendor_address:
        for ln in vendor_address.split("\n"):
            if ln.strip(): c.add_paragraph().add_run(ln)
    c.add_paragraph()
    p = c.add_paragraph(); p.add_run(f"JOB NO.: {job_number}").bold = True
    if job_location:
        c.add_paragraph().add_run(f"JOB LOCATION: {job_location}").bold = True

    c = row.cells[3]; c.merge(row.cells[5])
    p = c.paragraphs[0]; run = p.add_run("SHIP TO:"); run.bold = True; run.font.small_caps = True
    for ln in ship_to_lines: c.add_paragraph().add_run(ln)

    c = row.cells[6]; c.merge(row.cells[8])
    p = c.paragraphs[0]; run = p.add_run("P.O. Number:"); run.bold = True
    c.add_paragraph().add_run(po_number)
    c.add_paragraph().add_run(
        "The P.O. number must appear on all related correspondence, shipping papers, and invoices."
    ).font.size = Pt(7)

    # ── Spacer ──
    row = table.add_row(); row.cells[0].merge(row.cells[8])

    # ── PO metadata headers ──
    row = table.add_row()
    for i, h in enumerate(["P.O. DATE", "P.O. NUMBER", "REQUISITIONER",
                            "SHIPPED VIA", "F.O.B. POINT", "TERMS", "", "", ""]):
        if h:
            p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h); run.bold = True; run.font.size = Pt(7); run.font.small_caps = True
            _add_shading(row.cells[i], "D9E2F3")

    # ── PO metadata values ──
    row = table.add_row()
    for i, v in enumerate([po_date, po_number, requisitioner, shipped_via, fob_point, terms, "", "", ""]):
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(v)).font.size = Pt(9)

    # ── Spacer ──
    row = table.add_row(); row.cells[0].merge(row.cells[8])

    # ── Line item headers ──
    headers = ["ITEM#", "PROFILE\nNUMBER", "ALLOY/\nTEMPER", "FINISH",
               "LENGTH\n(INCHES)", "APPROX.\nLEAD TIME", "PCS",
               "UNIT\nPRICE/PC", "TOTAL"]
    row = table.add_row()
    for i, h in enumerate(headers):
        row.cells[i].width = col_widths[i]
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(7.5); run.font.small_caps = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _add_shading(row.cells[i], "2E75B6")

    # ── Line items ──
    for idx, line in enumerate(al_lines):
        row = table.add_row()
        line_total = line["unit_price"] * line["qty"]
        unit_price_str = f"${line['unit_price']:,.2f}" if line["unit_price"] else "-"
        tooling_str    = f"${line['tooling']:,.2f}" if line.get("tooling") and line["tooling"] else "-"
        vals = [
            str(idx + 1),
            line["profile"],
            line.get("alloy", "6063-T6"),
            line.get("finish", ""),
            str(line["length_in"]),
            line.get("lead_time", ""),
            str(line["qty"]),
            unit_price_str,
            f"${line_total:,.2f}" if line_total else "-",
        ]
        for i, v in enumerate(vals):
            row.cells[i].width = col_widths[i]
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 3: p.alignment = WD_ALIGN_PARAGRAPH.LEFT   # Finish left-aligned
            if i == 8: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Total right-aligned
            p.add_run(v).font.size = Pt(9)

    # ── Empty buffer rows ──
    for _ in range(3):
        row = table.add_row()
        row.cells[0].merge(row.cells[8])

    # ── Packing + Subtotal ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[5])
    run = c.paragraphs[0].add_run(f"Packing: {packaging_note}")
    run.underline = True; run.font.size = Pt(9)
    c = row.cells[6]; c.merge(row.cells[7])
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c.paragraphs[0].add_run("SUBTOTAL").font.size = Pt(9)
    row.cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[8].paragraphs[0].add_run(f"${subtotal:,.2f}"); run.bold = True; run.font.size = Pt(9)

    def _cost_row_al(label, amount_str):
        r = table.add_row()
        r.cells[0].merge(r.cells[5])
        # Instructions in left cell (first cost row only)
        c2 = r.cells[6]; c2.merge(r.cells[7])
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c2.paragraphs[0].add_run(label).font.size = Pt(9)
        r.cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r.cells[8].paragraphs[0].add_run(amount_str).font.size = Pt(9)

    _cost_row_al("SALES TAX", "")
    _cost_row_al("SHIPPING AND HANDLING", "")

    # ── Terms row ──
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[5])
    for ln in [
        "1. Enter this order in accordance with the prices, terms, delivery method, and specifications listed in this purchase order.",
        "2. Please notify us immediately if you are unable to ship as specified.",
        "3. Send Invoices to accounts@inovues.com",
        "4. Send all correspondence to:",
        "   INOVUES, INC.",
        "   2700 Post Oak Blvd, 2100, Houston, TX 77056",
        "   (833) 466-8837 (INO-VUES)",
        "   info@inovues.com",
    ]:
        c.add_paragraph().add_run(ln).font.size = Pt(7)
    c2 = row.cells[6]; c2.merge(row.cells[7])
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.paragraphs[0].add_run("TOTAL").bold = True
    row.cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[8].paragraphs[0].add_run(f"${subtotal:,.2f}"); run.bold = True; run.font.size = Pt(10)

    # ── Authorized by ──
    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    c = row.cells[5]; c.merge(row.cells[7])
    c.paragraphs[0].add_run("Authorized by _____________________").font.size = Pt(9)
    row.cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[8].paragraphs[0].add_run(po_date).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────
# PO DOCX GENERATOR (reused from SWR03111.py)
# ─────────────────────────────────────────────────────────────
def _add_shading(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shd  = tcPr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    tcPr.append(shd)

def generate_po_docx(
    vendor_name, vendor_contact, vendor_email, vendor_address,
    ship_to_lines, job_number, job_location,
    po_date, po_number, requisitioner,
    lead_time, shipped_via, fob_point, terms,
    glass_lines,
    price_per_sqft, packaging_cost, shipping_cost, sales_tax, other_cost,
    packaging_note="Non-returnable boxed crate/rack",
    logo_path=None,
    unit_label="ft²",
    price_label="Price/ft²",
):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin   = Cm(1.27)
        section.right_margin  = Cm(1.27)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    subtotal = sum(line["area_total"] * price_per_sqft for line in glass_lines)
    total    = subtotal + (sales_tax or 0) + (packaging_cost or 0) + (shipping_cost or 0) + (other_cost or 0)

    table = doc.add_table(rows=0, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    col_widths = [Cm(1.2), Cm(3.0), Cm(3.8), Cm(2.2), Cm(1.2), Cm(2.5), Cm(2.5)]
    for i, w in enumerate(col_widths):
        table.columns[i].width = w

    # Logo + PURCHASE ORDER header
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[2])
    if logo_path and os.path.exists(logo_path):
        c.paragraphs[0].add_run().add_picture(logo_path, width=Inches(1.5))
    c = row.cells[3]; c.merge(row.cells[6])
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("PURCHASE ORDER")
    run.bold = True; run.font.size = Pt(16); run.font.small_caps = True

    # TO / SHIP TO / BILL TO
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[1])
    p = c.paragraphs[0]; run = p.add_run("TO:"); run.bold = True; run.font.small_caps = True
    c.add_paragraph().add_run(vendor_name)
    if vendor_contact: c.add_paragraph().add_run(f"Attn: {vendor_contact}")
    if vendor_email:   c.add_paragraph().add_run(vendor_email)
    if vendor_address:
        for ln in vendor_address.split("\n"):
            if ln.strip(): c.add_paragraph().add_run(ln)
    c.add_paragraph()
    p = c.add_paragraph(); run = p.add_run(f"JOB NO.: {job_number}"); run.bold = True
    if job_location:
        p = c.add_paragraph(); run = p.add_run(f"JOB LOCATION: {job_location}"); run.bold = True

    c = row.cells[2]; c.merge(row.cells[4])
    p = c.paragraphs[0]; run = p.add_run("SHIP TO:"); run.bold = True; run.font.small_caps = True
    for ln in ship_to_lines: c.add_paragraph().add_run(ln)

    c = row.cells[5]; c.merge(row.cells[6])
    p = c.paragraphs[0]; run = p.add_run("BILL TO:"); run.bold = True; run.font.small_caps = True
    for ln in ["INOVUES, INC.", "2700 Post Oak Blvd., 2100", "Houston, TX 77056",
               "accounts@inovues.com", "(833) 466-8837 (INO-VUES)"]:
        c.add_paragraph().add_run(ln)

    # Spacer
    row = table.add_row(); row.cells[0].merge(row.cells[6])

    # PO metadata headers
    row = table.add_row()
    for i, h in enumerate(["P.O. DATE", "P.O. NUMBER", "REQUISITIONER", "LEAD TIME",
                            "SHIPPED VIA", "F.O.B. POINT", "TERMS"]):
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(7); run.font.small_caps = True
        _add_shading(row.cells[i], "D9E2F3")

    # PO metadata values
    row = table.add_row()
    for i, v in enumerate([po_date, po_number, requisitioner, lead_time, shipped_via, fob_point, terms]):
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(v)).font.size = Pt(9)

    # Spacer
    row = table.add_row(); row.cells[0].merge(row.cells[6])

    # Line item headers
    row = table.add_row()
    for i, h in enumerate(["ITEM#", "DESCRIPTION", f"UNIT SIZE", f"AREA EACH ({unit_label})",
                            "QTY", f"TOTAL ({unit_label})", "TOTAL"]):
        row.cells[i].width = col_widths[i]
        p = row.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(8); run.font.small_caps = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _add_shading(row.cells[i], "2E75B6")

    # Line items
    for idx, line in enumerate(glass_lines):
        row = table.add_row()
        line_total = line["area_total"] * price_per_sqft
        vals = [str(idx + 1), line.get("description", ""), line.get("size_str", ""),
                f"{line['area_each']:.3f}", str(line["qty"]),
                f"{line['area_total']:.3f}", f"${line_total:,.2f}"]
        for i, v in enumerate(vals):
            row.cells[i].width = col_widths[i]
            p = row.cells[i].paragraphs[0]
            if i == 0 or i >= 3: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 6:           p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(v).font.size = Pt(9)

    # Packaging + Subtotal
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[3])
    run = c.paragraphs[0].add_run(f"Packaging: {packaging_note}"); run.underline = True; run.font.size = Pt(9)
    c = row.cells[4]; c.merge(row.cells[5])
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c.paragraphs[0].add_run("SUBTOTAL").font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[6].paragraphs[0].add_run(f"${subtotal:,.2f}"); run.bold = True; run.font.size = Pt(9)

    def _cost_row(label, amount):
        r = table.add_row()
        r.cells[0].merge(r.cells[3])
        c2 = r.cells[4]; c2.merge(r.cells[5])
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c2.paragraphs[0].add_run(label).font.size = Pt(9)
        r.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r.cells[6].paragraphs[0].add_run(f"${amount:,.2f}" if amount else "").font.size = Pt(9)

    _cost_row("SALES TAX", sales_tax)
    _cost_row("PACKAGING", packaging_cost)

    # Shipping row
    row = table.add_row()
    c = row.cells[0]; c.merge(row.cells[3])
    for ln in [
        "1. Enter this order in accordance with the prices, terms, delivery method, and specifications listed in this purchase order.",
        "2. Please notify us immediately if you are unable to ship as specified.",
        "3. Send all correspondence to:",
        "   INOVUES, INC.",
        "   2700 Post Oak Blvd, 2100, Houston, TX 77056",
        "   (833) 466-8837 (INO-VUES)",
        "   accounts@inovues.com",
    ]:
        c.add_paragraph().add_run(ln).font.size = Pt(7)
    c2 = row.cells[4]; c2.merge(row.cells[5])
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.paragraphs[0].add_run("SHIPPING & HANDLING").font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].add_run(f"${shipping_cost:,.2f}" if shipping_cost else "").font.size = Pt(9)

    _cost_row("OTHER", other_cost)

    # Total
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    c = row.cells[4]; c.merge(row.cells[5])
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = c.paragraphs[0].add_run("TOTAL"); run.bold = True; run.font.size = Pt(10)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = row.cells[6].paragraphs[0].add_run(f"${total:,.2f}"); run.bold = True; run.font.size = Pt(10)

    # Signature
    row = table.add_row()
    row.cells[0].merge(row.cells[3])
    c = row.cells[4]; c.merge(row.cells[5])
    run = c.paragraphs[0].add_run("Authorized by _____________________"); run.italic = True; run.font.size = Pt(9)
    row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[6].paragraphs[0].add_run(po_date).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─────────────────────────────────────────────────────────────
# ODOO PO CREATION
# ─────────────────────────────────────────────────────────────
def create_odoo_po(vendor, po_lines, price_per_unit, project_number, project_name,
                    po_number, requisitioner, subtotal, grand_total, po_buf,
                    po_type, po_date_str,
                    payment_term_id=None, incoterm_id=None,
                    date_planned=None):
    """Create a draft purchase.order in Odoo.

    Newly-supported optional fields (Fix #2):
      payment_term_id — int, account.payment.term ID. Maps from the Terms dropdown.
      incoterm_id     — int, account.incoterms ID. Maps from the F.O.B. dropdown.
      date_planned    — datetime.date or datetime.datetime. Maps from the Lead Time date picker.

    The 'requisitioner' free-text field is fuzzy-matched against active internal
    users; if a match is found, user_id is set on the PO. Otherwise it remains
    free text in the .docx and chatter message only (preserves prior behavior).
    """
    uid, models = get_odoo_connection()

    def oc(model, method, args, kwargs={}):
        return odoo_call(models, uid, model, method, args, kwargs)

    # Find product
    product_name = "SWR Glass Panel" if po_type == "Glass" else "SWR Aluminium Extrusion"
    product_ids = oc("product.product", "search", [[("name", "ilike", product_name)]])
    if not product_ids:
        # Use any product as fallback
        product_ids = oc("product.product", "search", [[("active", "=", True)]], {"limit": 1})
    if not product_ids:
        raise Exception(f"No product found. Please create a product named '{product_name}' in Odoo.")

    product_id = product_ids[0]
    order_lines = []
    for line in po_lines:
        desc = (f"{line.get('description', '')}\n"
                f"{line.get('size_str', line.get('length_in', ''))}\n"
                f"Area/Length: {line['area_total']:.3f}  |  Qty: {line['qty']}")
        line_total = line["area_total"] * price_per_unit
        order_lines.append((0, 0, {
            "product_id":  product_id,
            "name":        desc,
            "product_qty": line["qty"],
            "price_unit":  line_total / line["qty"] if line["qty"] > 0 else 0.0,
        }))

    # Build the PO payload, only including optional fields when provided
    po_vals = {
        "partner_id":  vendor["id"],
        "origin":      f"{project_number} / {po_type} PO",
        "order_line":  order_lines,
    }
    if payment_term_id:
        po_vals["payment_term_id"] = payment_term_id
    if incoterm_id:
        po_vals["incoterm_id"] = incoterm_id
    if date_planned:
        # Odoo expects "YYYY-MM-DD HH:MM:SS"
        if hasattr(date_planned, "strftime"):
            po_vals["date_planned"] = date_planned.strftime("%Y-%m-%d 12:00:00") \
                if not hasattr(date_planned, "hour") else date_planned.strftime("%Y-%m-%d %H:%M:%S")
        else:
            po_vals["date_planned"] = str(date_planned)

    # Fuzzy-match requisitioner free-text to a real Odoo user
    user_map, _ = fetch_users()
    if requisitioner and user_map:
        matched_uid = user_map.get(requisitioner.strip().lower())
        if matched_uid:
            po_vals["user_id"] = matched_uid

    po_id = oc("purchase.order", "create", [po_vals])
    po_data = oc("purchase.order", "read", [po_id], {"fields": ["name"]})
    po_name = po_data[0]["name"] if po_data else f"ID {po_id}"

    # Attach docx
    oc("ir.attachment", "create", [{
        "name":      f"INOVUES_PO_{po_number}_{po_type}_{datetime.now().strftime('%Y%m%d')}.docx",
        "type":      "binary",
        "datas":     base64.b64encode(po_buf.getvalue()).decode("utf-8"),
        "res_model": "purchase.order",
        "res_id":    po_id,
        "mimetype":  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }])

    # Chatter message
    oc("purchase.order", "message_post", [[po_id]], {
        "body": (
            f"<b>📝 PO from INOVUES PO Generator</b><br/>"
            f"Project: {project_name}<br/>"
            f"PO Number: {po_number}<br/>"
            f"Requisitioner: {requisitioner}<br/>"
            f"Type: {po_type}<br/>"
            f"Vendor: {vendor['name']}<br/>"
            f"Subtotal: ${subtotal:,.2f} | Total: ${grand_total:,.2f}<br/>"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ),
        "message_type": "comment",
        "subtype_xmlid": "mail.mt_comment",
    })
    return po_id, po_name

# ─────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────
st.image("ilogo.png", width=180) if os.path.exists("ilogo.png") else None
st.title("📝 INOVUES PO Generator")
st.caption("Generate Glass and Aluminium Purchase Orders from Odoo project files")

# ── PO Type selector ──
po_type = st.radio("Select PO Type", ["Glass PO", "Aluminium PO"],
                   horizontal=True, label_visibility="collapsed")
po_type_key = "Glass" if po_type == "Glass PO" else "Aluminium"
file_keyword = "Glass" if po_type_key == "Glass" else "cutting_list"

st.divider()

# ── Step 1: Project + File selection ──
st.subheader("Step 1 — Select Project & File")

col_proj, col_file = st.columns([1, 2])

with col_proj:
    project_map, proj_err = fetch_projects()
    if proj_err:
        st.error(f"Could not load projects: {proj_err}")
        st.stop()

    selected_project = st.selectbox("Odoo Project", options=list(project_map.keys()),
                                    index=None, placeholder="Choose a project...")

attachments = []
selected_attachment = None
line_items = []

if selected_project:
    project_id = project_map[selected_project]
    attachments, att_err = fetch_project_attachments(project_id, po_type_key)

    if att_err:
        st.error(f"Could not fetch attachments: {att_err}")
    elif not attachments:
        st.warning("No Excel attachments found in this project's tasks.")
    else:
        # Only show files matching the file_keyword for this PO type
        relevant = [a for a in attachments if file_keyword.lower() in a["name"].lower()]

        if not relevant:
            with col_file:
                if po_type_key == "Glass":
                    st.warning(
                        "⚠️ **No Glass files found in this project.**\n\n"
                        "**Workflow to create one:**\n"
                        "1. Open the **SWR Cutlist** app\n"
                        "2. Run it on this project's window dimensions\n"
                        "3. Save to Odoo — it will attach `..._SWR_Glass_...xlsx` to an Engineering task\n"
                        "4. Come back here and refresh"
                    )
                else:
                    st.warning(
                        "⚠️ **No Cutting Optimizer files found in this project.**\n\n"
                        "**Workflow to create one:**\n"
                        "1. Open the **SWR Cutlist** app → save to Odoo (produces `..._SWR_AggCutOnly_...xlsx`)\n"
                        "2. Open the **Cutting Optimizer** app → load that AggCutOnly file → save to Odoo (produces `cutting_list_...xlsx`)\n"
                        "3. Come back here and refresh"
                    )
        else:
            with col_file:
                file_options = {f"{a['name']}  (task: {a['task_name']})": a for a in relevant}
                selected_file_key = st.selectbox(
                    f"Select {file_keyword} file",
                    options=list(file_options.keys()),
                    index=None,
                    placeholder="Choose a file..."
                )

            if selected_file_key:
                selected_attachment = file_options[selected_file_key]
                file_bytes = base64.b64decode(selected_attachment["datas"])

                if po_type_key == "Glass":
                    line_items, parse_err = parse_glass_file(file_bytes)
                else:
                    line_items, parse_err = parse_optimizer_file(file_bytes)

                if parse_err:
                    st.error(f"Could not parse file: {parse_err}")
                    line_items = []

# ── Manual upload fallback ──
upload_label = "cutting_list optimizer .xlsx file" if po_type_key == "Aluminium" else "Glass .xlsx file"
with st.expander("Or upload a file manually (fallback)", expanded=not selected_project):
    uploaded = st.file_uploader(f"Upload {upload_label}", type=["xlsx"])
    if uploaded:
        file_bytes = uploaded.read()
        if po_type_key == "Glass":
            line_items, parse_err = parse_glass_file(file_bytes)
        else:
            line_items, parse_err = parse_optimizer_file(file_bytes)
        if parse_err:
            st.error(f"Could not parse file: {parse_err}")
            line_items = []

# ══════════════════════════════════════════════════════════════
# GLASS PO FLOW
# ══════════════════════════════════════════════════════════════
if line_items and po_type_key == "Glass":
    st.divider()
    st.subheader(f"Step 2 — Review & Edit Line Items ({len(line_items)} lines)")

    total_qty  = sum(l["qty"] for l in line_items)
    total_area = sum(l["area_total"] for l in line_items)
    st.write(f"**{total_qty} pieces | {total_area:.2f} ft² total**")

    st.write("Enter description for each glass size:")
    first_desc = st.text_input("Description for all lines (default)",
                               placeholder="e.g. GT1 – 10mm Leadus VIG: 5Tlow-E+V+5T",
                               key="glass_desc_default")
    for line in line_items:
        line["description"] = first_desc

    preview = pd.DataFrame([{
        "Tag": l["tag"], "Size": l["size_str"],
        "Area Each (ft²)": l["area_each"], "Qty": l["qty"],
        "Area Total (ft²)": l["area_total"],
    } for l in line_items])
    st.dataframe(preview, use_container_width=True, hide_index=True)

    st.divider()
    st.subheader("Step 3 — Vendor & PO Details")

    vcol1, vcol2 = st.columns(2)
    with vcol1:
        vendor_map, vendor_err = fetch_vendors()
        if vendor_err:
            st.error(f"Could not load vendors: {vendor_err}")
        selected_vendor = st.selectbox("Select Vendor", options=list(vendor_map.keys()),
                                       index=None, placeholder="Choose a vendor...")
        if selected_vendor:
            v = vendor_map[selected_vendor]
            with st.expander("Vendor details", expanded=False):
                st.text(f"Email: {v.get('email','')}")
                st.text(f"Address: {v.get('full_address','')}")
    with vcol2:
        project_number = st.text_input("Project / PO Number", value="INO-")
        requisitioner  = st.text_input("Requisitioner", value="Stephan Ketterer")
        job_location   = st.text_input("Job Location", value="")

    dcol1, dcol2, dcol3 = st.columns(3)
    with dcol1:
        # Lead Time → date picker (writes to date_planned in Odoo)
        from datetime import timedelta
        lead_date = st.date_input("Expected Delivery Date",
                                  value=datetime.now().date() + timedelta(days=14),
                                  help="Becomes 'date_planned' on the Odoo PO.")
        lead_time = lead_date.strftime("%m/%d/%Y")  # human-readable for the .docx
        shipped_via = st.selectbox("Shipped Via", ["Air", "Ground"], index=1)
    with dcol2:
        # F.O.B. → dropdown of all 11 Odoo incoterms
        incoterm_map, incoterm_err = fetch_incoterms()
        if incoterm_err:
            st.warning(f"Could not load incoterms: {incoterm_err}")
        # Default to DDP (matches old hardcoded default)
        incoterm_keys = list(incoterm_map.keys()) if incoterm_map else []
        ddp_idx = next((i for i, k in enumerate(incoterm_keys) if k.startswith("DDP")), 0)
        fob_label = st.selectbox("F.O.B. Point", incoterm_keys,
                                  index=ddp_idx if incoterm_keys else None,
                                  help="Becomes 'incoterm_id' on the Odoo PO.")
        fob_point = fob_label.split(" — ")[0] if fob_label else ""  # short code for .docx

        # Terms → dropdown of payment terms from Odoo
        payterm_map, payterm_err = fetch_payment_terms()
        if payterm_err:
            st.warning(f"Could not load payment terms: {payterm_err}")
        payterm_keys = list(payterm_map.keys()) if payterm_map else []
        # Default to "30 Days" (closest match to old "Net 30" default)
        net30_idx = next((i for i, k in enumerate(payterm_keys) if "30" in k.lower() and "balance" not in k.lower() and "10th" not in k.lower()), 0)
        terms = st.selectbox("Payment Terms", payterm_keys,
                              index=net30_idx if payterm_keys else None,
                              help="Becomes 'payment_term_id' on the Odoo PO.")
    with dcol3:
        price_per_unit = st.number_input("Price per ft² ($)", value=0.0, min_value=0.0,
                                          step=0.01, format="%.2f")

    ship_to_text = st.text_area("Ship To", value=SHIP_TO_DEFAULT, height=100)

    pcol1, pcol2 = st.columns(2)
    with pcol1:
        packaging_cost = st.number_input("Packaging ($)", value=0.0, min_value=0.0, format="%.2f")
        sales_tax      = st.number_input("Sales Tax ($)", value=0.0, min_value=0.0, format="%.2f")
        packaging_note = st.text_input("Packaging Note", value="Non-returnable boxed crate/rack")
    with pcol2:
        shipping_cost = st.number_input("Shipping ($)", value=0.0, min_value=0.0, format="%.2f")
        other_cost    = st.number_input("Other ($)", value=0.0, min_value=0.0, format="%.2f")

    subtotal    = sum(l["area_total"] for l in line_items) * price_per_unit
    grand_total = subtotal + sales_tax + packaging_cost + shipping_cost + other_cost
    st.write(f"**Subtotal: ${subtotal:,.2f}  |  Grand Total: ${grand_total:,.2f}**")

    st.divider()
    st.subheader("Step 4 — Generate PO")
    vendor_ready = selected_vendor is not None and vendor_map
    gcol1, gcol2 = st.columns(2)

    with gcol1:
        if st.button("📄 Generate PO (.docx)", disabled=not vendor_ready, key="glass_gen"):
            v = vendor_map[selected_vendor]
            po_buf = generate_po_docx(
                vendor_name=v["name"], vendor_contact=v.get("contact_name", ""),
                vendor_email=v.get("email", ""), vendor_address=v.get("full_address", ""),
                ship_to_lines=[ln for ln in ship_to_text.split("\n") if ln.strip()],
                job_number=project_number, job_location=job_location,
                po_date=datetime.now().strftime("%m/%d/%Y"), po_number=project_number,
                requisitioner=requisitioner, lead_time=lead_time, shipped_via=shipped_via,
                fob_point=fob_point, terms=terms, glass_lines=line_items,
                price_per_sqft=price_per_unit, packaging_cost=packaging_cost,
                shipping_cost=shipping_cost, sales_tax=sales_tax, other_cost=other_cost,
                packaging_note=packaging_note, unit_label="ft²",
            )
            st.session_state["po_buf"] = po_buf.getvalue()
            st.session_state["po_buf_name"] = f"INOVUES_PO_{project_number}_Glass_{datetime.now().strftime('%Y%m%d')}.docx"
            st.success("✅ PO document ready!")
        if "po_buf" in st.session_state:
            st.download_button("💾 Download PO .docx", data=st.session_state["po_buf"],
                               file_name=st.session_state.get("po_buf_name", "INOVUES_PO.docx"),
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with gcol2:
        if st.button("📝 Create Draft PO in Odoo", type="primary", disabled=not vendor_ready, key="glass_odoo"):
            if "po_buf" not in st.session_state:
                st.warning("Generate the .docx first.")
            else:
                with st.spinner("Creating PO in Odoo..."):
                    try:
                        v = vendor_map[selected_vendor]
                        po_id, po_name = create_odoo_po(
                            vendor=v, po_lines=line_items, price_per_unit=price_per_unit,
                            project_number=project_number, project_name=selected_project or "",
                            po_number=project_number, requisitioner=requisitioner,
                            subtotal=subtotal, grand_total=grand_total,
                            po_buf=io.BytesIO(st.session_state["po_buf"]),
                            po_type="Glass", po_date_str=datetime.now().strftime("%m/%d/%Y"),
                            payment_term_id=payterm_map.get(terms) if terms else None,
                            incoterm_id=incoterm_map.get(fob_label) if fob_label else None,
                            date_planned=lead_date,
                        )
                        st.success(f"✅ Draft PO **{po_name}** created in Odoo!")
                        st.info(f"🔗 {ODOO_URL}/web#id={po_id}&model=purchase.order&view_type=form")
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

# ══════════════════════════════════════════════════════════════
# ALUMINIUM PO FLOW
# ══════════════════════════════════════════════════════════════
elif line_items and po_type_key == "Aluminium":
    st.divider()
    st.subheader(f"Step 2 — Review & Edit Line Items ({len(line_items)} lines)")
    st.caption("Pre-filled from optimizer output. Adjust quantities and fill in pricing details.")

    # ── Per-line editable fields ──
    st.write("**Edit each line item:**")

    # Shared fields header
    sh1, sh2, sh3, sh4, sh5, sh6, sh7 = st.columns([1.5, 1.5, 2, 2, 1, 1.5, 1.5])
    sh1.markdown("**Profile #**"); sh2.markdown("**Length (in)**")
    sh3.markdown("**Alloy/Temper**"); sh4.markdown("**Finish**")
    sh5.markdown("**Qty**"); sh6.markdown("**Lead Time**"); sh7.markdown("**Unit Price ($)**")

    for i, line in enumerate(line_items):
        c1, c2, c3, c4, c5, c6, c7 = st.columns([1.5, 1.5, 2, 2, 1, 1.5, 1.5])
        line["profile"]    = c1.text_input("", value=line["profile"],    key=f"al_prof_{i}", label_visibility="collapsed")
        line["length_in"]  = c2.number_input("", value=float(line["length_in"]), min_value=0.0, step=1.0, key=f"al_len_{i}", label_visibility="collapsed")
        line["alloy"]      = c3.text_input("", value=line.get("alloy", "6063-T6"), key=f"al_alloy_{i}", label_visibility="collapsed")
        line["finish"]     = c4.text_input("", value=line.get("finish", ""),       key=f"al_finish_{i}", label_visibility="collapsed")
        line["qty"]        = c5.number_input("", value=int(line["qty"]), min_value=0, step=1, key=f"al_qty_{i}", label_visibility="collapsed")
        line["lead_time"]  = c6.text_input("", value=line.get("lead_time", "3-4 weeks"), key=f"al_lead_{i}", label_visibility="collapsed")
        line["unit_price"] = c7.number_input("", value=float(line.get("unit_price", 0.0)), min_value=0.0, step=0.01, format="%.2f", key=f"al_price_{i}", label_visibility="collapsed")

    total_pcs = sum(l["qty"] for l in line_items)
    subtotal_al = sum(l["unit_price"] * l["qty"] for l in line_items)
    st.write(f"**{total_pcs} total pcs | Subtotal: ${subtotal_al:,.2f}**")

    st.divider()
    st.subheader("Step 3 — Vendor & PO Details")

    vcol1, vcol2 = st.columns(2)
    with vcol1:
        vendor_map, vendor_err = fetch_vendors()
        if vendor_err:
            st.error(f"Could not load vendors: {vendor_err}")
        selected_vendor = st.selectbox("Select Vendor", options=list(vendor_map.keys()),
                                       index=None, placeholder="Choose a vendor...")
        if selected_vendor:
            v = vendor_map[selected_vendor]
            with st.expander("Vendor details", expanded=False):
                st.text(f"Email: {v.get('email','')}")
                st.text(f"Address: {v.get('full_address','')}")
    with vcol2:
        project_number = st.text_input("Project / PO Number", value="INO-", key="al_proj_num")
        requisitioner  = st.text_input("Requisitioner", value="Stephan Ketterer", key="al_req")
        job_location   = st.text_input("Job Location", value="", key="al_job_loc")

    dcol1, dcol2 = st.columns(2)
    with dcol1:
        from datetime import timedelta
        lead_date_al = st.date_input("Expected Delivery Date",
                                     value=datetime.now().date() + timedelta(days=28),
                                     key="al_lead_date",
                                     help="Becomes 'date_planned' on the Odoo PO.")
        lead_time_al = lead_date_al.strftime("%m/%d/%Y")
        shipped_via = st.selectbox("Shipped Via", ["Air", "Ground"], index=1, key="al_ship")
        # F.O.B. dropdown
        incoterm_map, incoterm_err = fetch_incoterms()
        if incoterm_err:
            st.warning(f"Could not load incoterms: {incoterm_err}")
        incoterm_keys = list(incoterm_map.keys()) if incoterm_map else []
        # Aluminium previously defaulted to "Shipping Point" (not a real incoterm).
        # Use FOB as the closest official equivalent.
        fob_idx = next((i for i, k in enumerate(incoterm_keys) if k.startswith("FOB")), 0)
        fob_label = st.selectbox("F.O.B. Point", incoterm_keys,
                                  index=fob_idx if incoterm_keys else None,
                                  key="al_fob",
                                  help="Becomes 'incoterm_id' on the Odoo PO.")
        fob_point = fob_label.split(" — ")[0] if fob_label else ""
    with dcol2:
        # Payment terms dropdown
        payterm_map, payterm_err = fetch_payment_terms()
        if payterm_err:
            st.warning(f"Could not load payment terms: {payterm_err}")
        payterm_keys = list(payterm_map.keys()) if payterm_map else []
        # Aluminium default was "Per Quote" — closest match in Odoo is "30 Days"
        net30_idx = next((i for i, k in enumerate(payterm_keys) if "30" in k.lower() and "balance" not in k.lower() and "10th" not in k.lower()), 0)
        terms = st.selectbox("Payment Terms", payterm_keys,
                              index=net30_idx if payterm_keys else None,
                              key="al_terms",
                              help="Becomes 'payment_term_id' on the Odoo PO.")
        packaging_note  = st.text_input("Packing Note", value="Fully Corrugated Bundles, Paper Layer Separation", key="al_pack_note")

    ship_to_text = st.text_area("Ship To", value=SHIP_TO_DEFAULT, height=100, key="al_ship_to")

    grand_total_al = subtotal_al
    st.write(f"**Grand Total: ${grand_total_al:,.2f}**")

    st.divider()
    st.subheader("Step 4 — Generate PO")
    vendor_ready = selected_vendor is not None and vendor_map
    gcol1, gcol2 = st.columns(2)

    with gcol1:
        if st.button("📄 Generate Aluminium PO (.docx)", disabled=not vendor_ready, key="al_gen"):
            v = vendor_map[selected_vendor]
            po_buf = generate_aluminum_po_docx(
                vendor_name=v["name"], vendor_contact=v.get("contact_name", ""),
                vendor_email=v.get("email", ""), vendor_address=v.get("full_address", ""),
                ship_to_lines=[ln for ln in ship_to_text.split("\n") if ln.strip()],
                job_number=project_number, job_location=job_location,
                po_date=datetime.now().strftime("%m/%d/%Y"), po_number=project_number,
                requisitioner=requisitioner, shipped_via=shipped_via,
                fob_point=fob_point, terms=terms,
                al_lines=line_items,
                packaging_note=packaging_note,
            )
            st.session_state["al_po_buf"] = po_buf.getvalue()
            st.session_state["al_po_buf_name"] = f"INOVUES_PO_{project_number}_Aluminium_{datetime.now().strftime('%Y%m%d')}.docx"
            st.success("✅ Aluminium PO document ready!")
        if "al_po_buf" in st.session_state:
            st.download_button("💾 Download Aluminium PO .docx",
                               data=st.session_state["al_po_buf"],
                               file_name=st.session_state.get("al_po_buf_name", "INOVUES_PO_Aluminium.docx"),
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key="al_dl")

    with gcol2:
        if st.button("📝 Create Draft PO in Odoo", type="primary", disabled=not vendor_ready, key="al_odoo"):
            if "al_po_buf" not in st.session_state:
                st.warning("Generate the .docx first.")
            else:
                with st.spinner("Creating PO in Odoo..."):
                    try:
                        v = vendor_map[selected_vendor]
                        # Map al_lines to the format create_odoo_po expects
                        odoo_lines = [{
                            "description": f"{l['profile']} — {l['length_in']}\" extrusion",
                            "size_str":    f"{l['length_in']}\"",
                            "area_total":  float(l["qty"]),
                            "qty":         l["qty"],
                        } for l in line_items]
                        po_id, po_name = create_odoo_po(
                            vendor=v, po_lines=odoo_lines, price_per_unit=1.0,
                            project_number=project_number, project_name=selected_project or "",
                            po_number=project_number, requisitioner=requisitioner,
                            subtotal=subtotal_al, grand_total=grand_total_al,
                            po_buf=io.BytesIO(st.session_state["al_po_buf"]),
                            po_type="Aluminium", po_date_str=datetime.now().strftime("%m/%d/%Y"),
                            payment_term_id=payterm_map.get(terms) if terms else None,
                            incoterm_id=incoterm_map.get(fob_label) if fob_label else None,
                            date_planned=lead_date_al,
                        )
                        st.success(f"✅ Draft PO **{po_name}** created in Odoo!")
                        st.info(f"🔗 {ODOO_URL}/web#id={po_id}&model=purchase.order&view_type=form")
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")

else:
    if selected_project:
        st.info("Select a file above to continue.")
    else:
        st.info("Select an Odoo project to get started.")
